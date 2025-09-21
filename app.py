# app.py
from flask import Flask, request, jsonify, send_file, abort
from docx import Document
import os

app = Flask(__name__)

# file we always write to
FILE_NAME = "Summary of the Sale of Goodwill.docx"

# Optional: use a secret token for safety (set SECRET_TOKEN in Render env vars)
SECRET_TOKEN = os.environ.get("SECRET_TOKEN", None)

def check_auth():
    if SECRET_TOKEN:
        auth = request.headers.get("Authorization", "")
        # Expect: Authorization: Bearer <token>
        if not auth.startswith("Bearer ") or auth.split(" ", 1)[1] != SECRET_TOKEN:
            abort(401)

@app.route("/write_word", methods=["POST"])
def write_word():
    # optional auth
    check_auth()

    data = request.get_json(force=True)
    content = data.get("content", "")

    if os.path.exists(FILE_NAME):
        doc = Document(FILE_NAME)
    else:
        doc = Document()

    # Split into lines and add each as paragraph
    for line in content.split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line)

    doc.save(FILE_NAME)
    return jsonify({"message": f"Written to {FILE_NAME}", "content_lines": len(content.splitlines())})

# Optional: download the file from server
@app.route("/download_word", methods=["GET"])
def download_word():
    check_auth()
    if os.path.exists(FILE_NAME):
        return send_file(FILE_NAME, as_attachment=True, download_name=FILE_NAME)
    return jsonify({"error": "File not found"}), 404

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
